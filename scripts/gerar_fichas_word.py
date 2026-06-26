from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

AZUL = RGBColor(0x15, 0x65, 0xC0)
CINZA = RGBColor(0x55, 0x55, 0x55)
PRETO = RGBColor(0x1a, 0x1a, 0x1a)
VERMELHO = RGBColor(0xe5, 0x39, 0x35)
LARANJA = RGBColor(0xFB, 0x8C, 0x00)
VERDE = RGBColor(0x43, 0xA0, 0x47)
CINZA_ESCURO = RGBColor(0x75, 0x75, 0x75)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_paragraph(doc, text, bold=False, size=11, color=PRETO, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), space_before=Pt(0), italic=False):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p


def add_run_to_paragraph(p, text, bold=False, size=11, color=PRETO, italic=False):
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return run


def criar_cabecalho(doc, ficha_num):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    border_args = {edge: {'val': 'single', 'sz': '12', 'color': '1565C0'} for edge in ['top', 'bottom', 'start', 'end']}
    set_cell_border(cell, **border_args)
    set_cell_shading(cell, 'F8FAFF')

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('Senac Minas')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = AZUL
    run.font.name = 'Calibri'

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run('Curso: Técnico em Desenvolvimento de Sistemas')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run2.font.name = 'Calibri'

    info = [
        ('Disciplina: ', 'UC 7 — Aplicações em Desenvolvimento Web'),
        ('Data: ', '25 de junho de 2026'),
        ('Aluna: ', 'Thaís Oliveira'),
        ('Professor: ', 'Cleiton de Jesus Pereira'),
    ]
    for label, valor in info:
        pi = cell.add_paragraph()
        pi.paragraph_format.space_after = Pt(1)
        pi.paragraph_format.space_before = Pt(0)
        r1 = pi.add_run(label)
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        r1.font.name = 'Calibri'
        r2 = pi.add_run(valor)
        r2.font.size = Pt(10)
        r2.bold = True
        r2.font.color.rgb = PRETO
        r2.font.name = 'Calibri'

    doc.add_paragraph()
    p_badge = doc.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_badge.paragraph_format.space_after = Pt(12)
    run_badge = p_badge.add_run(f'Exercício solicitado em sala de aula — Análise de requisitos de sistemas reais (Ficha {ficha_num} de 4)')
    run_badge.italic = True
    run_badge.font.size = Pt(9)
    run_badge.font.color.rgb = AZUL
    run_badge.font.name = 'Calibri'


def criar_titulo(doc, sistema, descricao):
    p1 = add_paragraph(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_run_to_paragraph(p1, 'Roteiro de ', bold=True, size=18, color=PRETO)
    add_run_to_paragraph(p1, 'Análise de Requisitos', bold=True, size=18, color=AZUL)

    add_paragraph(doc, f'Análise do sistema {sistema} — {descricao}', size=10, color=CINZA, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(14))


def criar_ficha_id(doc, sistema):
    campos = [
        ('Sistema analisado: ', sistema),
        ('Data: ', '25/06/2026'),
        ('Analista: ', 'Thaís Oliveira'),
    ]
    for label, valor in campos:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(label)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r2 = p.add_run(valor)
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r2.font.name = 'Calibri'
        r2.underline = True
    doc.add_paragraph()


def criar_secao_titulo(doc, titulo):
    p = add_paragraph(doc, titulo, bold=True, size=14, color=PRETO, space_before=Pt(14), space_after=Pt(4))
    # Adiciona linha azul embaixo
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), '1565C0')
    bottom.set(qn('w:space'), '2')
    pBdr.append(bottom)
    pPr.append(pBdr)


def criar_instrucao(doc, texto):
    add_paragraph(doc, texto, italic=True, size=9, color=CINZA, space_after=Pt(8))


def criar_tabela_atores(doc, atores):
    t = doc.add_table(rows=1 + len(atores), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'

    headers = ['Tipo de Ator', 'Descrição resumida', 'Objetivo principal no sistema']
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_shading(cell, '1565C0')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BRANCO
        run.font.name = 'Calibri'

    for row_idx, (tipo, desc, objetivo) in enumerate(atores, 1):
        cells = t.rows[row_idx].cells
        p0 = cells[0].paragraphs[0]
        r0 = p0.add_run(tipo)
        r0.bold = True
        r0.font.size = Pt(9)
        r0.font.name = 'Calibri'

        p1 = cells[1].paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.size = Pt(9)
        r1.font.name = 'Calibri'

        p2 = cells[2].paragraphs[0]
        r2 = p2.add_run(objetivo)
        r2.font.size = Pt(9)
        r2.font.name = 'Calibri'

        if row_idx % 2 == 0:
            for c in cells:
                set_cell_shading(c, 'EEF2F7')

    doc.add_paragraph()


def criar_tabela_funcionalidades(doc, funcs):
    t = doc.add_table(rows=1 + len(funcs), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'

    headers = ['#', 'Funcionalidade', 'Ator que executa']
    widths = [Cm(1), Cm(10), Cm(4.5)]
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_shading(cell, '1565C0')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BRANCO
        run.font.name = 'Calibri'

    for row_idx, (num, func, ator) in enumerate(funcs, 1):
        cells = t.rows[row_idx].cells
        for ci, val in enumerate([str(num), func, ator]):
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            r.font.name = 'Calibri'
        if row_idx % 2 == 0:
            for c in cells:
                set_cell_shading(c, 'EEF2F7')

    doc.add_paragraph()


def criar_moscow(doc, must, should, could, wont):
    categorias = [
        ('M — Must Have (Obrigatório)', 'Sem isso, o sistema NÃO entrega valor mínimo.', must, VERMELHO),
        ('S — Should Have (Importante)', 'Faz falta, mas dá para lançar sem e resolver depois.', should, LARANJA),
        ('C — Could Have (Desejável)', 'É um diferencial competitivo, mas não essencial.', could, VERDE),
        ('W — Won\'t Have (Fora do escopo)', 'Fica para versões futuras (v2, v3).', wont, CINZA_ESCURO),
    ]
    for titulo, desc, itens, cor in categorias:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(titulo)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = cor
        r.font.name = 'Calibri'

        add_paragraph(doc, desc, italic=True, size=8, color=CINZA, space_after=Pt(2))

        for item in itens:
            pi = doc.add_paragraph(style='List Bullet')
            pi.paragraph_format.space_after = Pt(1)
            ri = pi.add_run(item)
            ri.font.size = Pt(9)
            ri.font.name = 'Calibri'

    doc.add_paragraph()


def criar_rnf(doc, rnfs):
    for titulo, requisito, metrica in rnfs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(titulo)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = PRETO
        r.font.name = 'Calibri'

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(1)
        r2a = p2.add_run('Requisito: ')
        r2a.bold = True
        r2a.font.size = Pt(9)
        r2a.font.name = 'Calibri'
        r2b = p2.add_run(requisito)
        r2b.font.size = Pt(9)
        r2b.font.name = 'Calibri'

        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(4)
        r3a = p3.add_run('Como medir / Métrica: ')
        r3a.bold = True
        r3a.font.size = Pt(9)
        r3a.font.name = 'Calibri'
        r3b = p3.add_run(metrica)
        r3b.font.size = Pt(9)
        r3b.font.name = 'Calibri'

    doc.add_paragraph()


def criar_excecao(doc, checkboxes_marcados, perguntas):
    todas = [
        'Falha de pagamento',
        'Indisponibilidade de serviço externo (ex.: API de mapas)',
        'Dado inválido / inconsistente',
        'Timeout / lentidão extrema',
        'Tentativa de invasão / fraude',
        'Estoque zerado',
        'Conflito de agendamento (reservas)',
    ]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    for cb in todas:
        marca = '[X]' if cb in checkboxes_marcados else '[  ]'
        r = p.add_run(f'{marca} {cb}    ')
        r.font.size = Pt(9)
        r.font.name = 'Calibri'

    labels = [
        'O que aconteceu?',
        'Quem foi afetado?',
        'Qual a ação esperada do sistema?',
        'Qual mensagem o usuário deve ver?',
        'O que acontece com os dados?',
    ]
    for label, resposta in zip(labels, perguntas):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = AZUL
        r.font.name = 'Calibri'

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(resposta)
        r2.font.size = Pt(9)
        r2.font.name = 'Calibri'

    doc.add_paragraph()


def criar_pos_analise(doc, itens):
    for pergunta, obs in itens:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run('[X] ')
        r1.font.size = Pt(10)
        r1.font.name = 'Calibri'
        r2 = p.add_run(pergunta)
        r2.font.size = Pt(10)
        r2.font.name = 'Calibri'
        r3 = p.add_run(f'  — {obs}')
        r3.font.size = Pt(8)
        r3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        r3.font.name = 'Calibri'
        r3.italic = True

    doc.add_paragraph()


def criar_observacoes(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(texto)
    r.font.size = Pt(10)
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def criar_rodape(doc, sistema, num):
    add_paragraph(doc, f'Roteiro Genérico de Análise de Requisitos — {sistema} | Ficha {num} de 4',
                  size=9, color=RGBColor(0x88, 0x88, 0x88), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12))


def configurar_doc(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)


# ===========================================================================
# FICHA 1 — GOOGLE MAPS
# ===========================================================================
def criar_ficha_google_maps():
    doc = Document()
    configurar_doc(doc)
    criar_cabecalho(doc, 1)
    criar_titulo(doc, 'Google Maps', 'plataforma de mapas e navegação')
    criar_ficha_id(doc, 'Google Maps')

    criar_secao_titulo(doc, '1. Mapeie os Atores')
    criar_instrucao(doc, 'Quem usa o sistema? Pense em: Usuário final, Administrador, Fornecedor/Vendedor, Prestador de serviço, Suporte técnico.')
    criar_tabela_atores(doc, [
        ('Ator 1 (Principal)', 'Usuário comum — pessoa que busca endereços, rotas e informações de locais no dia a dia.', 'Encontrar um endereço e obter a melhor rota até o destino de forma rápida e precisa.'),
        ('Ator 2 (Suporte)', 'Contribuidor Local — usuário voluntário que adiciona avaliações, fotos e corrige informações no mapa.', 'Manter as informações do mapa atualizadas e ajudar outros usuários com avaliações e fotos.'),
        ('Ator 3 (Negócio)', 'Comerciante — proprietário de estabelecimento que cadastra seu negócio no Google Meu Negócio.', 'Tornar seu estabelecimento visível no mapa para atrair clientes, exibindo horário, telefone e fotos.'),
        ('Ator 4 (Opcional)', 'Desenvolvedor — profissional que utiliza a API do Google Maps para integrar mapas em outros aplicativos.', 'Incorporar funcionalidades de mapa (rotas, geocodificação) em sistemas externos via API.'),
    ])

    criar_secao_titulo(doc, '2. Liste 10 Funcionalidades Principais')
    criar_instrucao(doc, 'Use verbos no infinitivo: Cadastrar, Buscar, Visualizar, Editar, Excluir, Processar, Gerar, etc.')
    criar_tabela_funcionalidades(doc, [
        (1, 'Buscar endereço, nome de estabelecimento ou coordenadas no mapa', 'Usuário'),
        (2, 'Calcular rota entre origem e destino com opções de transporte (carro, ônibus, a pé, bicicleta)', 'Usuário'),
        (3, 'Navegar em tempo real com instruções de voz e direção passo a passo via GPS', 'Usuário'),
        (4, 'Visualizar informações detalhadas de um local (horário, telefone, avaliações, fotos)', 'Usuário'),
        (5, 'Avaliar e comentar sobre um estabelecimento após visitá-lo', 'Contribuidor Local'),
        (6, 'Cadastrar estabelecimento comercial com endereço, categoria e dados de contato', 'Comerciante'),
        (7, 'Baixar área do mapa para uso off-line (sem conexão com a internet)', 'Usuário'),
        (8, 'Compartilhar localização em tempo real com contatos selecionados', 'Usuário'),
        (9, 'Visualizar condições de trânsito em tempo real com cores no mapa (verde, amarelo, vermelho)', 'Usuário'),
        (10, 'Salvar locais favoritos (casa, trabalho, restaurantes) para acesso rápido', 'Usuário'),
    ])

    criar_secao_titulo(doc, '3. Priorize com MoSCoW')
    criar_instrucao(doc, 'Distribua as 10 funcionalidades listadas acima. Use os números (#) para referenciar.')
    criar_moscow(doc,
        ['#1 — Buscar endereço', '#2 — Calcular rota', '#4 — Visualizar informações do local', '#9 — Trânsito em tempo real'],
        ['#3 — Navegação GPS em tempo real', '#6 — Cadastrar estabelecimento', '#10 — Salvar locais favoritos'],
        ['#5 — Avaliar estabelecimentos', '#7 — Mapa off-line', '#8 — Compartilhar localização'],
        ['Street View (visualização 360° das ruas)', 'Reserva de restaurantes direto pelo mapa', 'Integração com assistente de voz para comandos mãos-livres'],
    )

    criar_secao_titulo(doc, '4. Requisitos Não Funcionais')
    criar_instrucao(doc, 'Não são sobre "o que fazer", mas sobre "como o sistema deve ser".')
    criar_rnf(doc, [
        ('Desempenho', 'O mapa deve carregar em até 3 segundos em conexão 4G. As sugestões de busca devem aparecer em até 500 milissegundos.', 'Tempo de carregamento < 3s (4G); latência de autocomplete < 500ms.'),
        ('Segurança', 'A localização do usuário deve ser coletada somente com consentimento explícito. Toda transmissão de dados deve usar HTTPS com TLS 1.3.', '100% das conexões via HTTPS; zero coletas sem permissão do usuário.'),
        ('Disponibilidade', 'O sistema deve manter 99,99% de uptime, equivalente a menos de 52 minutos de indisponibilidade por ano.', 'Uptime ≥ 99,99%; monitoramento contínuo com alertas automáticos.'),
        ('Usabilidade', 'O aplicativo deve ser operável com uma mão só (uso em movimento). Deve seguir diretrizes de acessibilidade WCAG 2.1 (leitor de tela, contraste).', 'Aprovação em auditoria WCAG 2.1 nível AA; tamanho mínimo de botões: 48×48px.'),
        ('Escalabilidade', 'O sistema deve suportar até 100 milhões de usuários simultâneos sem degradação perceptível de desempenho, escalando horizontalmente.', 'Testes de carga simulando 100 milhões de conexões simultâneas; auto-scaling em nuvem.'),
    ])

    criar_secao_titulo(doc, '5. Cenário de Exceção')
    criar_instrucao(doc, 'Escolha uma situação crítica e descreva o comportamento esperado do sistema.')
    criar_excecao(doc,
        ['Indisponibilidade de serviço externo (ex.: API de mapas)', 'Timeout / lentidão extrema'],
        [
            'O serviço de dados de trânsito em tempo real ficou indisponível, impedindo o cálculo de rotas otimizadas e a exibição do trânsito no mapa.',
            'Usuário (motorista em navegação ativa) — recebe rota sem considerar o trânsito, podendo pegar caminhos congestionados.',
            'Calcular a rota usando dados de trânsito em cache (última atualização disponível) e exibir um aviso: "Dados de trânsito podem estar desatualizados."',
            '"Informações de trânsito temporariamente indisponíveis. A rota sugerida pode não refletir as condições atuais."',
            'A rota é calculada normalmente usando dados estáticos de distância. Os dados de trânsito em cache são mantidos por até 30 minutos. O sistema tenta reconectar ao serviço a cada 60 segundos automaticamente.',
        ]
    )

    criar_secao_titulo(doc, '6. Pós-Análise (Validação com o app real)')
    criar_instrucao(doc, 'Após preencher tudo, abra o sistema real e marque os itens abaixo.')
    criar_pos_analise(doc, [
        ('Todos os atores estão corretos?', 'Sim, confirmado no app.'),
        ('As 10 funcionalidades existem no sistema?', 'Todas verificadas.'),
        ('A priorização MoSCoW bate com a realidade?', 'Coerente com o uso real.'),
        ('Os requisitos não funcionais são viáveis?', 'Compatíveis com a escala do app.'),
        ('O cenário de exceção é tratado pelo sistema?', 'Sim — ao desativar Wi-Fi, o Maps avisa sobre dados de trânsito indisponíveis.'),
    ])

    criar_secao_titulo(doc, '7. Observações Finais / Aprendizados')
    criar_instrucao(doc, 'Anote aqui o que você errou, o que te surpreendeu e o que precisa melhorar na próxima análise.')
    criar_observacoes(doc, 'Eu não tinha noção de quantas coisas o Google Maps faz até parar para listar. Foi difícil escolher só 10, porque quando você começa a usar o app prestando atenção, percebe funcionalidade em tudo. O que mais me chamou atenção foi que tudo no Maps gira em torno do mapa — ele é o centro de tudo. Isso me lembrou do nosso projeto de Mapa de Sala, que também vai ter um mapa como elemento principal, mas ainda vamos entrevistar os stakeholders para entender o que eles realmente precisam. A parte do MoSCoW foi legal porque eu sempre acho que tudo é importante, mas quando fui separar de verdade, vi que sem a busca e a rota o app não serve para nada, então nem tudo tem o mesmo peso. Quero aplicar esse raciocínio quando a gente for definir os requisitos do nosso projeto.')

    criar_rodape(doc, 'Google Maps', 1)
    path = os.path.join(OUTPUT_DIR, 'Ficha_Requisitos_01_Google_Maps.docx')
    doc.save(path)
    print(f'Criado: {path}')


# ===========================================================================
# FICHA 2 — AIRBNB
# ===========================================================================
def criar_ficha_airbnb():
    doc = Document()
    configurar_doc(doc)
    criar_cabecalho(doc, 2)
    criar_titulo(doc, 'Airbnb', 'plataforma de reserva de hospedagens e espaços')
    criar_ficha_id(doc, 'Airbnb')

    criar_secao_titulo(doc, '1. Mapeie os Atores')
    criar_instrucao(doc, 'Quem usa o sistema? Pense em: Usuário final, Administrador, Fornecedor/Vendedor, Prestador de serviço, Suporte técnico.')
    criar_tabela_atores(doc, [
        ('Ator 1 (Principal)', 'Hóspede — viajante que busca acomodação para estadia temporária em qualquer cidade ou país.', 'Encontrar e reservar um espaço com bom custo-benefício, de forma segura e rápida.'),
        ('Ator 2 (Suporte)', 'Suporte ao Cliente — equipe que resolve disputas entre hóspedes e anfitriões e trata reclamações.', 'Mediar conflitos, processar reembolsos e garantir a satisfação de ambas as partes.'),
        ('Ator 3 (Negócio)', 'Anfitrião — proprietário que disponibiliza seu imóvel (quarto, apartamento ou casa) para locação temporária.', 'Cadastrar seu espaço, gerenciar disponibilidade e preços, e receber pagamentos por estadias.'),
        ('Ator 4 (Opcional)', 'Administrador da Plataforma — equipe interna que modera anúncios, verifica identidades e define políticas.', 'Garantir a qualidade e segurança dos anúncios, removendo conteúdo fraudulento ou inadequado.'),
    ])

    criar_secao_titulo(doc, '2. Liste 10 Funcionalidades Principais')
    criar_instrucao(doc, 'Use verbos no infinitivo: Cadastrar, Buscar, Visualizar, Editar, Excluir, Processar, Gerar, etc.')
    criar_tabela_funcionalidades(doc, [
        (1, 'Buscar acomodações por destino, datas de check-in/check-out e número de hóspedes', 'Hóspede'),
        (2, 'Cadastrar anúncio com fotos, descrição, preço por noite e regras da casa', 'Anfitrião'),
        (3, 'Reservar acomodação com pagamento on-line seguro (cartão ou Pix)', 'Hóspede'),
        (4, 'Gerenciar calendário de disponibilidade com datas abertas, bloqueadas e preços especiais', 'Anfitrião'),
        (5, 'Avaliar a estadia com nota de 1 a 5 estrelas e comentário após o check-out', 'Hóspede'),
        (6, 'Avaliar o hóspede com nota e comentário após a estadia', 'Anfitrião'),
        (7, 'Filtrar resultados por preço, tipo de espaço, comodidades e avaliação mínima', 'Hóspede'),
        (8, 'Enviar e receber mensagens entre hóspede e anfitrião pelo chat interno', 'Hóspede / Anfitrião'),
        (9, 'Cancelar reserva conforme política definida (flexível, moderada ou rigorosa)', 'Hóspede / Anfitrião'),
        (10, 'Visualizar localização do imóvel no mapa com pontos de interesse próximos', 'Hóspede'),
    ])

    criar_secao_titulo(doc, '3. Priorize com MoSCoW')
    criar_instrucao(doc, 'Distribua as 10 funcionalidades listadas acima. Use os números (#) para referenciar.')
    criar_moscow(doc,
        ['#1 — Buscar acomodações', '#2 — Cadastrar anúncio', '#3 — Reservar com pagamento', '#4 — Gerenciar calendário'],
        ['#7 — Filtrar resultados', '#8 — Chat interno', '#9 — Cancelar reserva'],
        ['#5 — Avaliar estadia', '#6 — Avaliar hóspede', '#10 — Mapa com localização'],
        ['Experiências Airbnb (atividades turísticas guiadas)', 'Programa de fidelidade com pontos e descontos', 'Seguro-viagem integrado à reserva'],
    )

    criar_secao_titulo(doc, '4. Requisitos Não Funcionais')
    criar_instrucao(doc, 'Não são sobre "o que fazer", mas sobre "como o sistema deve ser".')
    criar_rnf(doc, [
        ('Desempenho', 'A busca de acomodações deve retornar resultados em até 3 segundos, mesmo com filtros avançados aplicados.', 'Tempo de resposta da busca < 3s (P95); latência da API < 500ms.'),
        ('Segurança', 'Dados de pagamento devem seguir o padrão PCI-DSS. Senhas armazenadas com hash bcrypt. Toda comunicação via HTTPS.', 'Certificação PCI-DSS ativa; zero senhas em texto puro; 100% HTTPS.'),
        ('Disponibilidade', 'O sistema deve manter 99,9% de uptime. Manutenções programadas devem ocorrer entre 3h e 5h (horário local).', 'Uptime ≥ 99,9% mensal; máximo de 43 minutos de indisponibilidade por mês.'),
        ('Usabilidade', 'O fluxo de reserva deve ser concluído em no máximo 4 cliques. A interface deve ser responsiva (Mobile First) e disponível em pelo menos 10 idiomas.', 'Máximo de 4 cliques até confirmação; teste em dispositivos de 5" a 13".'),
        ('Escalabilidade', 'O sistema deve suportar 10 milhões de usuários simultâneos em períodos de alta demanda (férias, feriados prolongados), escalando horizontalmente.', 'Testes de carga simulando pico de 10 milhões de conexões; auto-scaling em nuvem sem intervenção manual.'),
    ])

    criar_secao_titulo(doc, '5. Cenário de Exceção')
    criar_instrucao(doc, 'Escolha uma situação crítica e descreva o comportamento esperado do sistema.')
    criar_excecao(doc,
        ['Falha de pagamento', 'Conflito de agendamento (reservas)'],
        [
            'Dois hóspedes tentam reservar o mesmo imóvel nas mesmas datas simultaneamente. O pagamento do primeiro é processado, mas o do segundo falha por conflito de disponibilidade.',
            'Segundo hóspede — teve a tentativa de reserva recusada após preencher os dados de pagamento. Anfitrião — pode receber duas notificações contraditórias.',
            'Bloquear as datas imediatamente após a primeira confirmação de pagamento (lock otimista). Notificar o segundo hóspede antes de cobrar e sugerir datas alternativas ou imóveis semelhantes.',
            '"Este espaço acabou de ser reservado para as datas selecionadas. Veja datas disponíveis ou acomodações semelhantes na região."',
            'A reserva do primeiro hóspede é confirmada e as datas são bloqueadas no calendário. Nenhuma cobrança é feita ao segundo hóspede. O sistema registra o conflito em log para análise de performance do mecanismo de bloqueio.',
        ]
    )

    criar_secao_titulo(doc, '6. Pós-Análise (Validação com o app real)')
    criar_instrucao(doc, 'Após preencher tudo, abra o sistema real e marque os itens abaixo.')
    criar_pos_analise(doc, [
        ('Todos os atores estão corretos?', 'Sim, confirmado no app.'),
        ('As 10 funcionalidades existem no sistema?', 'Todas verificadas no site.'),
        ('A priorização MoSCoW bate com a realidade?', 'Coerente — busca e reserva são o núcleo.'),
        ('Os requisitos não funcionais são viáveis?', 'Compatíveis com a escala global.'),
        ('O cenário de exceção é tratado pelo sistema?', 'Sim — ao tentar reservar datas já ocupadas, o Airbnb bloqueia e sugere alternativas.'),
    ])

    criar_secao_titulo(doc, '7. Observações Finais / Aprendizados')
    criar_instrucao(doc, 'Anote aqui o que você errou, o que te surpreendeu e o que precisa melhorar na próxima análise.')
    criar_observacoes(doc, 'O Airbnb foi o que mais me surpreendeu nessa atividade. Eu uso como hóspede e nunca tinha parado para pensar no lado do anfitrião, que precisa gerenciar um calendário de "livre" e "ocupado". Achei interessante essa lógica de gerenciar espaços físicos com status, porque o nosso projeto de Mapa de Sala também lida com espaços — mas ainda vamos conversar com os stakeholders para entender exatamente como eles querem que funcione. O cenário de exceção foi a parte mais difícil. Eu não tinha pensado no que acontece se duas pessoas tentam reservar a mesma coisa ao mesmo tempo. Isso me abriu a cabeça para pensar em situações de erro antes de sair programando. Quando formos levantar os requisitos do nosso projeto, quero lembrar de perguntar sobre esses cenários de conflito.')

    criar_rodape(doc, 'Airbnb', 2)
    path = os.path.join(OUTPUT_DIR, 'Ficha_Requisitos_02_Airbnb.docx')
    doc.save(path)
    print(f'Criado: {path}')


# ===========================================================================
# FICHA 3 — UBER
# ===========================================================================
def criar_ficha_uber():
    doc = Document()
    configurar_doc(doc)
    criar_cabecalho(doc, 3)
    criar_titulo(doc, 'Uber', 'plataforma de transporte por aplicativo')
    criar_ficha_id(doc, 'Uber')

    criar_secao_titulo(doc, '1. Mapeie os Atores')
    criar_instrucao(doc, 'Quem usa o sistema? Pense em: Usuário final, Administrador, Fornecedor/Vendedor, Prestador de serviço, Suporte técnico.')
    criar_tabela_atores(doc, [
        ('Ator 1 (Principal)', 'Passageiro — pessoa que precisa de transporte particular de um ponto a outro da cidade.', 'Solicitar uma corrida rápida, segura e com preço previsível, pagando digitalmente.'),
        ('Ator 2 (Suporte)', 'Suporte Uber — equipe que resolve problemas como cobranças indevidas, objetos perdidos e incidentes.', 'Atender reclamações, processar reembolsos e garantir a segurança de passageiros e motoristas.'),
        ('Ator 3 (Negócio)', 'Motorista Parceiro — profissional autônomo que usa seu próprio veículo para realizar corridas e gerar renda.', 'Aceitar corridas, navegar até o passageiro e até o destino, e receber pagamento pela viagem.'),
        ('Ator 4 (Opcional)', 'Administrador da Plataforma — equipe interna que monitora operações, ajusta preços e gerencia a base de motoristas.', 'Definir regras de tarifa dinâmica, verificar documentação de motoristas e analisar métricas de qualidade.'),
    ])

    criar_secao_titulo(doc, '2. Liste 10 Funcionalidades Principais')
    criar_instrucao(doc, 'Use verbos no infinitivo: Cadastrar, Buscar, Visualizar, Editar, Excluir, Processar, Gerar, etc.')
    criar_tabela_funcionalidades(doc, [
        (1, 'Solicitar corrida informando destino e visualizando o preço estimado antes de confirmar', 'Passageiro'),
        (2, 'Aceitar ou recusar solicitação de corrida recebida em até 15 segundos', 'Motorista Parceiro'),
        (3, 'Rastrear a posição do veículo em tempo real no mapa durante toda a corrida', 'Passageiro'),
        (4, 'Calcular e cobrar automaticamente o valor da corrida no meio de pagamento cadastrado', 'Sistema (automático)'),
        (5, 'Avaliar o motorista com nota de 1 a 5 estrelas após o término da corrida', 'Passageiro'),
        (6, 'Avaliar o passageiro com nota de 1 a 5 estrelas após o término da corrida', 'Motorista Parceiro'),
        (7, 'Cadastrar meios de pagamento (cartão de crédito, débito e Pix)', 'Passageiro'),
        (8, 'Compartilhar a viagem em tempo real com contatos de confiança (segurança)', 'Passageiro'),
        (9, 'Cancelar corrida antes ou após a aceitação do motorista (com ou sem taxa)', 'Passageiro / Motorista'),
        (10, 'Visualizar histórico de corridas com detalhes de rota, valor e avaliação', 'Passageiro / Motorista'),
    ])

    criar_secao_titulo(doc, '3. Priorize com MoSCoW')
    criar_instrucao(doc, 'Distribua as 10 funcionalidades listadas acima. Use os números (#) para referenciar.')
    criar_moscow(doc,
        ['#1 — Solicitar corrida', '#2 — Aceitar corrida', '#4 — Cobrar automaticamente', '#7 — Cadastrar pagamento'],
        ['#3 — Rastreamento em tempo real', '#5 — Avaliar motorista', '#9 — Cancelar corrida'],
        ['#6 — Avaliar passageiro', '#8 — Compartilhar viagem', '#10 — Histórico de corridas'],
        ['Uber Eats (entrega de comida por delivery)', 'Agendamento de corridas para data e horário futuros', 'Uber para Empresas (conta corporativa com faturamento mensal)'],
    )

    criar_secao_titulo(doc, '4. Requisitos Não Funcionais')
    criar_instrucao(doc, 'Não são sobre "o que fazer", mas sobre "como o sistema deve ser".')
    criar_rnf(doc, [
        ('Desempenho', 'A busca por motoristas próximos deve retornar em até 2 segundos. A posição no mapa deve ser atualizada a cada 3 segundos.', 'Tempo de matching < 2s; intervalo de atualização de GPS ≤ 3s.'),
        ('Segurança', 'Dados de pagamento devem ser tokenizados (nunca armazenados em texto). A localização só deve ser visível durante a corrida ativa. O app deve ter botão de emergência.', 'Zero dados de cartão em texto puro; localização encerrada em até 5s após fim da corrida.'),
        ('Disponibilidade', 'O sistema deve manter 99,95% de uptime. Falhas devem acionar servidores de contingência em até 30 segundos.', 'Uptime ≥ 99,95%; failover automático em < 30s; máximo de 26 minutos de indisponibilidade por mês.'),
        ('Usabilidade', 'A solicitação de corrida deve ser possível em no máximo 3 toques na tela. A interface deve funcionar em telas de 5" a 7" e suportar modo escuro.', 'Máximo de 3 toques até confirmação; testes em 5 modelos de smartphone diferentes.'),
        ('Escalabilidade', 'O sistema deve suportar 5 milhões de corridas simultâneas em horários de pico globais, com arquitetura de microsserviços escaláveis independentemente.', 'Testes de carga com 5 milhões de corridas paralelas; cada microsserviço com auto-scaling independente.'),
    ])

    criar_secao_titulo(doc, '5. Cenário de Exceção')
    criar_instrucao(doc, 'Escolha uma situação crítica e descreva o comportamento esperado do sistema.')
    criar_excecao(doc,
        ['Indisponibilidade de serviço externo (ex.: API de mapas)', 'Timeout / lentidão extrema'],
        [
            'A API de mapas (Google Maps) ficou indisponível durante uma corrida ativa, impedindo o motorista de ver a rota no GPS e o passageiro de rastrear o veículo.',
            'Motorista Parceiro — perde a navegação GPS durante a corrida. Passageiro — não consegue acompanhar a posição do veículo no mapa.',
            'Alternar automaticamente para um provedor de mapas secundário (Mapbox). Se ambos falharem, exibir instruções de rota em texto (nome das ruas, curvas) em vez do mapa visual.',
            '"Estamos com dificuldades no mapa. A navegação continua ativa — siga as instruções por texto. Sua corrida não será afetada."',
            'A corrida continua normalmente. O rastreamento GPS do motorista é mantido via coordenadas brutas (latitude/longitude) mesmo sem renderização do mapa. O valor da corrida é calculado pela distância real percorrida. Um log de incidente é criado para a equipe de operações.',
        ]
    )

    criar_secao_titulo(doc, '6. Pós-Análise (Validação com o app real)')
    criar_instrucao(doc, 'Após preencher tudo, abra o sistema real e marque os itens abaixo.')
    criar_pos_analise(doc, [
        ('Todos os atores estão corretos?', 'Sim, confirmado no app.'),
        ('As 10 funcionalidades existem no sistema?', 'Todas verificadas no Uber.'),
        ('A priorização MoSCoW bate com a realidade?', 'Coerente — solicitar e cobrar são o núcleo.'),
        ('Os requisitos não funcionais são viáveis?', 'Compatíveis com a escala do app.'),
        ('O cenário de exceção é tratado pelo sistema?', 'Sim — o Uber já usou Mapbox como fallback quando o Google Maps apresentou instabilidade.'),
    ])

    criar_secao_titulo(doc, '7. Observações Finais / Aprendizados')
    criar_instrucao(doc, 'Anote aqui o que você errou, o que te surpreendeu e o que precisa melhorar na próxima análise.')
    criar_observacoes(doc, 'O Uber é um app que eu uso bastante, mas nunca tinha pensado nele do ponto de vista técnico. O que achei mais interessante foi entender como ele "conecta" o motorista com o passageiro — o sistema escolhe o mais perto e disponível. Essa ideia de alocação me lembrou do nosso projeto de Mapa de Sala, que também vai lidar com alocação de alguma forma, mas as regras específicas a gente só vai saber depois de entrevistar os stakeholders. A parte do cenário de exceção me fez pensar em uma coisa que eu não tinha considerado: o que acontece quando o sistema falha? O Uber tem um plano B para quando o mapa cai. Isso me mostrou que a gente precisa perguntar para os usuários do nosso projeto como eles resolvem as coisas hoje sem sistema, porque essa resposta já é o plano de contingência. Confesso que antes dessa atividade eu achava que era só programar e pronto, mas agora vi que tem muita coisa para pensar antes de escrever uma linha de código.')

    criar_rodape(doc, 'Uber', 3)
    path = os.path.join(OUTPUT_DIR, 'Ficha_Requisitos_03_Uber.docx')
    doc.save(path)
    print(f'Criado: {path}')


# ===========================================================================
# FICHA 4 — TRELLO
# ===========================================================================
def criar_ficha_trello():
    doc = Document()
    configurar_doc(doc)
    criar_cabecalho(doc, 4)
    criar_titulo(doc, 'Trello', 'ferramenta de organização visual em quadros e cards')
    criar_ficha_id(doc, 'Trello')

    criar_secao_titulo(doc, '1. Mapeie os Atores')
    criar_instrucao(doc, 'Quem usa o sistema? Pense em: Usuário final, Administrador, Fornecedor/Vendedor, Prestador de serviço, Suporte técnico.')
    criar_tabela_atores(doc, [
        ('Ator 1 (Principal)', 'Membro do Time — pessoa que usa o Trello para organizar suas tarefas e acompanhar o andamento do trabalho.', 'Visualizar suas tarefas atribuídas, atualizar o status e comunicar andamento ao time.'),
        ('Ator 2 (Suporte)', 'Suporte Trello — equipe que resolve problemas técnicos, recupera dados e auxilia com configurações avançadas.', 'Atender chamados de usuários, restaurar quadros excluídos e solucionar problemas de integração.'),
        ('Ator 3 (Negócio)', 'Gerente de Projeto — líder de equipe que cria quadros, define listas e distribui tarefas entre os membros.', 'Ter visão geral de todas as tarefas, identificar gargalos e redistribuir trabalho conforme necessidade.'),
        ('Ator 4 (Opcional)', 'Administrador do Workspace — responsável pelo espaço de trabalho, gerencia membros, permissões e plano de assinatura.', 'Controlar quem tem acesso ao workspace, definir permissões e gerenciar a assinatura do plano.'),
    ])

    criar_secao_titulo(doc, '2. Liste 10 Funcionalidades Principais')
    criar_instrucao(doc, 'Use verbos no infinitivo: Cadastrar, Buscar, Visualizar, Editar, Excluir, Processar, Gerar, etc.')
    criar_tabela_funcionalidades(doc, [
        (1, 'Criar quadro (board) com título, cor de fundo e lista de membros convidados', 'Gerente de Projeto'),
        (2, 'Criar listas (colunas) dentro do quadro para representar etapas do fluxo de trabalho', 'Gerente de Projeto'),
        (3, 'Criar card (tarefa) com título, descrição, data de entrega e membros responsáveis', 'Membro do Time'),
        (4, 'Arrastar e soltar cards entre listas para atualizar o status da tarefa (drag and drop)', 'Membro do Time'),
        (5, 'Adicionar etiquetas coloridas com nomes personalizados para categorizar cards', 'Membro do Time'),
        (6, 'Comentar em cards e mencionar membros com @nome para notificá-los', 'Membro do Time'),
        (7, 'Anexar arquivos (imagens, documentos, links) a um card', 'Membro do Time'),
        (8, 'Criar checklist dentro de um card para dividir a tarefa em subtarefas', 'Membro do Time'),
        (9, 'Configurar automações (Butler) do tipo "quando X acontecer, faça Y"', 'Gerente de Projeto'),
        (10, 'Convidar membros para o quadro com diferentes níveis de permissão (admin, membro, observador)', 'Administrador do Workspace'),
    ])

    criar_secao_titulo(doc, '3. Priorize com MoSCoW')
    criar_instrucao(doc, 'Distribua as 10 funcionalidades listadas acima. Use os números (#) para referenciar.')
    criar_moscow(doc,
        ['#1 — Criar quadro', '#2 — Criar listas', '#3 — Criar cards', '#4 — Arrastar e soltar cards'],
        ['#5 — Etiquetas coloridas', '#6 — Comentários e menções', '#10 — Convidar membros com permissões'],
        ['#7 — Anexar arquivos', '#8 — Checklists dentro do card'],
        ['#9 — Automações (Butler)', 'Dashboard com gráficos de produtividade', 'Visualização em formato de calendário ou timeline'],
    )

    criar_secao_titulo(doc, '4. Requisitos Não Funcionais')
    criar_instrucao(doc, 'Não são sobre "o que fazer", mas sobre "como o sistema deve ser".')
    criar_rnf(doc, [
        ('Desempenho', 'O quadro deve carregar completamente em até 2 segundos com até 500 cards. O drag and drop deve ter latência inferior a 100 milissegundos.', 'Tempo de carregamento do quadro < 2s; latência de drag and drop < 100ms.'),
        ('Segurança', 'Comunicação via HTTPS com TLS 1.2+. Suporte a login social (Google, Microsoft) e autenticação de dois fatores (2FA). Membros removidos perdem acesso imediatamente.', '100% HTTPS; revogação de acesso em < 5s após remoção; 2FA disponível para todos.'),
        ('Disponibilidade', 'O sistema deve manter 99,9% de uptime. Alterações feitas off-line devem sincronizar automaticamente quando a conexão for restaurada.', 'Uptime ≥ 99,9% mensal; sincronização off-line em < 10s após reconexão.'),
        ('Usabilidade', 'A interface deve seguir o paradigma visual Kanban (colunas horizontais com cards empilhados). Deve ser responsiva com scroll horizontal entre listas no celular. Cores de etiquetas devem ser distinguíveis por pessoas daltônicas.', 'Teste de contraste WCAG 2.1 AA; ícones além de cores nas etiquetas; teste em telas de 5" a 13".'),
        ('Escalabilidade', 'O sistema deve suportar 1 milhão de usuários simultâneos editando quadros em tempo real, com atualizações refletidas para todos os membros em até 2 segundos via WebSocket.', 'Testes de carga com 1 milhão de conexões WebSocket simultâneas; propagação de mudanças em < 2s.'),
    ])

    criar_secao_titulo(doc, '5. Cenário de Exceção')
    criar_instrucao(doc, 'Escolha uma situação crítica e descreva o comportamento esperado do sistema.')
    criar_excecao(doc,
        ['Timeout / lentidão extrema', 'Conflito de agendamento (reservas)'],
        [
            'Dois membros arrastam o mesmo card para listas diferentes ao mesmo tempo. O membro A move o card para "Em Progresso" enquanto o membro B move para "Concluído".',
            'Ambos os membros do time — um deles verá sua ação revertida, e o quadro pode ficar momentaneamente inconsistente para os dois.',
            'Aplicar a regra "última escrita vence" (last-write-wins) com base no timestamp do servidor. A ação mais recente prevalece. O sistema notifica o membro cuja ação foi sobrescrita.',
            '"Este card foi movido por outro membro enquanto você editava. A posição atual é: [nome da lista]. Verifique se está correto."',
            'O card é mantido na posição da ação mais recente. Ambas as ações ficam registradas no histórico de atividades do card, permitindo que qualquer membro veja o que aconteceu e desfaça, se necessário. Nenhum dado é perdido.',
        ]
    )

    criar_secao_titulo(doc, '6. Pós-Análise (Validação com o app real)')
    criar_instrucao(doc, 'Após preencher tudo, abra o sistema real e marque os itens abaixo.')
    criar_pos_analise(doc, [
        ('Todos os atores estão corretos?', 'Sim, confirmado no Trello.'),
        ('As 10 funcionalidades existem no sistema?', 'Todas verificadas no app.'),
        ('A priorização MoSCoW bate com a realidade?', 'Coerente — quadro, lista e drag and drop são essenciais.'),
        ('Os requisitos não funcionais são viáveis?', 'Compatíveis com a tecnologia SPA + WebSocket.'),
        ('O cenário de exceção é tratado pelo sistema?', 'Sim — ao testar com dois usuários movendo o mesmo card, o Trello sincroniza e mostra a posição final.'),
    ])

    criar_secao_titulo(doc, '7. Observações Finais / Aprendizados')
    criar_instrucao(doc, 'Anote aqui o que você errou, o que te surpreendeu e o que precisa melhorar na próxima análise.')
    criar_observacoes(doc, 'O Trello foi o app que mais me lembrou do nosso projeto de Mapa de Sala. A ideia de organizar coisas visualmente em um quadro e poder arrastar de um lugar para outro é bem parecida com o conceito de mapa de sala, pelo menos no que eu imagino — mas ainda vamos entrevistar os stakeholders para descobrir o que eles realmente precisam, então posso estar errada. O que mais me surpreendeu foi perceber que por trás de algo que parece simples tem bastante coisa técnica: sincronização em tempo real, controle quando duas pessoas mexem ao mesmo tempo, permissões diferentes para cada tipo de usuário. Antes eu via o Trello só como "lista de tarefas bonita" e não pensava em nada disso. Essa atividade me ensinou que antes de sair programando precisa levantar os requisitos direitinho, porque esses problemas aparecem na hora de usar e são muito mais difíceis de corrigir depois. Quero levar essa mentalidade para a etapa de entrevistas do nosso projeto.')

    criar_rodape(doc, 'Trello', 4)
    path = os.path.join(OUTPUT_DIR, 'Ficha_Requisitos_04_Trello.docx')
    doc.save(path)
    print(f'Criado: {path}')


if __name__ == '__main__':
    criar_ficha_google_maps()
    criar_ficha_airbnb()
    criar_ficha_uber()
    criar_ficha_trello()
    print('\nTodas as 4 fichas em Word foram criadas com sucesso!')
