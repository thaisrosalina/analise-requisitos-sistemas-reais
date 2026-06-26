from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

AZUL = RGBColor(0x15, 0x65, 0xC0)
CINZA = RGBColor(0x55, 0x55, 0x55)
PRETO = RGBColor(0x1a, 0x1a, 0x1a)
VERMELHO = RGBColor(0xe5, 0x39, 0x35)
LARANJA = RGBColor(0xFB, 0x8C, 0x00)
VERDE = RGBColor(0x43, 0xA0, 0x47)
CINZA_ESCURO = RGBColor(0x75, 0x75, 0x75)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
CINZA_FONTE = RGBColor(0x88, 0x88, 0x88)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_paragraph(doc, text, bold=False, size=11, color=PRETO, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), space_before=Pt(0), italic=False):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p


def add_run_to_paragraph(p, text, bold=False, size=11, color=PRETO, italic=False):
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return run


def criar_cabecalho(doc, ficha_num):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    border_args = {edge: {'val': 'single', 'sz': '12', 'color': '1565C0'} for edge in ['top', 'bottom', 'start', 'end']}
    set_cell_border(cell, **border_args)
    set_cell_shading(cell, 'F8FAFF')

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('Senac Minas')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = AZUL
    run.font.name = 'Calibri'

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run('Curso: Técnico em Desenvolvimento de Sistemas')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run2.font.name = 'Calibri'

    info = [
        ('Disciplina: ', 'UC 7: Aplicações em Desenvolvimento Web'),
        ('Data: ', '25 de junho de 2026'),
        ('Aluna: ', 'Thaís Oliveira'),
        ('Professor: ', 'Cleiton de Jesus Pereira'),
    ]
    for label, valor in info:
        pi = cell.add_paragraph()
        pi.paragraph_format.space_after = Pt(1)
        pi.paragraph_format.space_before = Pt(0)
        r1 = pi.add_run(label)
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        r1.font.name = 'Calibri'
        r2 = pi.add_run(valor)
        r2.font.size = Pt(10)
        r2.bold = True
        r2.font.color.rgb = PRETO
        r2.font.name = 'Calibri'

    doc.add_paragraph()
    p_badge = doc.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_badge.paragraph_format.space_after = Pt(12)
    run_badge = p_badge.add_run(f'Exercício solicitado em sala de aula: Análise de requisitos de sistemas reais (Ficha {ficha_num} de 4). Versão com análise aprofundada')
    run_badge.italic = True
    run_badge.font.size = Pt(9)
    run_badge.font.color.rgb = AZUL
    run_badge.font.name = 'Calibri'


def criar_titulo(doc, sistema, descricao):
    p1 = add_paragraph(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_run_to_paragraph(p1, 'Roteiro de ', bold=True, size=18, color=PRETO)
    add_run_to_paragraph(p1, 'Análise de Requisitos', bold=True, size=18, color=AZUL)
    add_paragraph(doc, f'Análise do sistema {sistema}: {descricao}', size=10, color=CINZA, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(14))


def criar_ficha_id(doc, sistema):
    campos = [
        ('Sistema analisado: ', sistema),
        ('Data: ', '25/06/2026'),
        ('Analista: ', 'Thaís Oliveira'),
    ]
    for label, valor in campos:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(label)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r2 = p.add_run(valor)
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r2.font.name = 'Calibri'
        r2.underline = True
    doc.add_paragraph()


def criar_secao_titulo(doc, titulo):
    p = add_paragraph(doc, titulo, bold=True, size=14, color=PRETO, space_before=Pt(14), space_after=Pt(4))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), '1565C0')
    bottom.set(qn('w:space'), '2')
    pBdr.append(bottom)
    pPr.append(pBdr)


def criar_instrucao(doc, texto):
    add_paragraph(doc, texto, italic=True, size=9, color=CINZA, space_after=Pt(8))


def criar_fonte(doc, texto):
    add_paragraph(doc, texto, italic=True, size=8, color=CINZA_FONTE, space_after=Pt(4))


def criar_tabela_atores(doc, atores):
    t = doc.add_table(rows=1 + len(atores), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    headers = ['Tipo de Ator', 'Descrição resumida', 'Objetivo principal no sistema']
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_shading(cell, '1565C0')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BRANCO
        run.font.name = 'Calibri'
    for row_idx, (tipo, desc, objetivo) in enumerate(atores, 1):
        cells = t.rows[row_idx].cells
        for ci, val in enumerate([tipo, desc, objetivo]):
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.bold = (ci == 0)
            r.font.size = Pt(9)
            r.font.name = 'Calibri'
        if row_idx % 2 == 0:
            for c in cells:
                set_cell_shading(c, 'EEF2F7')
    doc.add_paragraph()


def criar_tabela_funcionalidades(doc, funcs):
    t = doc.add_table(rows=1 + len(funcs), cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    headers = ['#', 'Funcionalidade', 'Ator que executa']
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_shading(cell, '1565C0')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BRANCO
        run.font.name = 'Calibri'
    for row_idx, (num, func, ator) in enumerate(funcs, 1):
        cells = t.rows[row_idx].cells
        for ci, val in enumerate([str(num), func, ator]):
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            r.font.name = 'Calibri'
        if row_idx % 2 == 0:
            for c in cells:
                set_cell_shading(c, 'EEF2F7')
    doc.add_paragraph()


def criar_moscow(doc, must, should, could, wont, nota_analitica=''):
    categorias = [
        ('M - Must Have (Obrigatório)', 'Sem isso, o sistema NÃO entrega valor mínimo.', must, VERMELHO),
        ('S - Should Have (Importante)', 'Faz falta, mas dá para lançar sem e resolver depois.', should, LARANJA),
        ('C - Could Have (Desejável)', 'Diferencial competitivo, não essencial.', could, VERDE),
        ("W - Won't Have (Fora do escopo)", 'Fica para versões futuras.', wont, CINZA_ESCURO),
    ]
    for titulo, desc, itens, cor in categorias:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(titulo)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = cor
        r.font.name = 'Calibri'
        add_paragraph(doc, desc, italic=True, size=8, color=CINZA, space_after=Pt(2))
        for item in itens:
            pi = doc.add_paragraph(style='List Bullet')
            pi.paragraph_format.space_after = Pt(1)
            ri = pi.add_run(item)
            ri.font.size = Pt(9)
            ri.font.name = 'Calibri'
    if nota_analitica:
        add_paragraph(doc, nota_analitica, italic=True, size=8, color=CINZA_FONTE, space_after=Pt(4), space_before=Pt(6))
    doc.add_paragraph()


def criar_rnf(doc, rnfs):
    for titulo, requisito, metrica in rnfs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(titulo)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = PRETO
        r.font.name = 'Calibri'
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(1)
        r2a = p2.add_run('Requisito: ')
        r2a.bold = True
        r2a.font.size = Pt(9)
        r2a.font.name = 'Calibri'
        r2b = p2.add_run(requisito)
        r2b.font.size = Pt(9)
        r2b.font.name = 'Calibri'
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(4)
        r3a = p3.add_run('Como medir / Métrica: ')
        r3a.bold = True
        r3a.font.size = Pt(9)
        r3a.font.name = 'Calibri'
        r3b = p3.add_run(metrica)
        r3b.font.size = Pt(9)
        r3b.font.name = 'Calibri'
    doc.add_paragraph()


def criar_excecao(doc, checkboxes_marcados, perguntas):
    todas = [
        'Falha de pagamento',
        'Indisponibilidade de serviço externo',
        'Dado inválido / inconsistente',
        'Timeout / lentidão extrema',
        'Tentativa de invasão / fraude',
        'Estoque zerado',
        'Conflito de agendamento',
        'Falha de GPS / perda de sinal',
        'Nenhum motorista disponível',
        'Edição simultânea / conflito de dados',
        'Perda de conexão WebSocket',
    ]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    for cb in todas:
        marca = '[X]' if cb in checkboxes_marcados else '[  ]'
        r = p.add_run(f'{marca} {cb}    ')
        r.font.size = Pt(9)
        r.font.name = 'Calibri'
    labels = [
        'O que aconteceu?',
        'Quem foi afetado?',
        'Qual a ação esperada do sistema?',
        'Qual mensagem o usuário deve ver?',
        'O que acontece com os dados?',
    ]
    for label, resposta in zip(labels, perguntas):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = AZUL
        r.font.name = 'Calibri'
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(resposta)
        r2.font.size = Pt(9)
        r2.font.name = 'Calibri'
    doc.add_paragraph()


def criar_pos_analise(doc, itens):
    for pergunta, obs in itens:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run('[X] ')
        r1.font.size = Pt(10)
        r1.font.name = 'Calibri'
        r2 = p.add_run(pergunta)
        r2.font.size = Pt(10)
        r2.font.name = 'Calibri'
        r3 = p.add_run(f'  - {obs}')
        r3.font.size = Pt(8)
        r3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        r3.font.name = 'Calibri'
        r3.italic = True
    doc.add_paragraph()


def criar_observacoes(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(texto)
    r.font.size = Pt(10)
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def criar_rodape(doc, sistema, num):
    add_paragraph(doc, f'Roteiro de Análise de Requisitos: {sistema} | Ficha {num} de 4 | Análise Aprofundada com Dados Reais',
                  size=9, color=CINZA_FONTE, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(12))


def configurar_doc(doc):
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)


# ===========================================================================
# FICHA 1 — GOOGLE MAPS (Análise Aprofundada)
# ===========================================================================
def criar_ficha_google_maps():
    doc = Document()
    configurar_doc(doc)
    criar_cabecalho(doc, 1)
    criar_titulo(doc, 'Google Maps', 'plataforma de mapeamento e navegação, 2 bilhões de usuários ativos mensais')
    criar_ficha_id(doc, 'Google Maps')

    criar_secao_titulo(doc, '1. Mapeie os Atores')
    criar_instrucao(doc, 'Atores enriquecidos com dados reais de escala e reclassificação de papéis baseada em análise de impacto.')
    criar_tabela_atores(doc, [
        ('Ator 1 (Principal)', 'Usuário final, com 2 bilhões de MAU (Monthly Active Users). Consome dados de mapa, rotas e informações de locais.', 'Encontrar endereço e obter a melhor rota até o destino de forma rápida e precisa.'),
        ('Ator 2 (Colaborador)', 'Local Guide, programa de contribuição voluntária com 150+ milhões de participantes em 24.000+ cidades. Reclassificado de "Suporte" para "Colaborador" por gerar conteúdo que alimenta o sistema.', 'Manter informações do mapa atualizadas contribuindo avaliações, fotos e correções.'),
        ('Ator 3 (Negócio)', 'Comerciante, com 250+ milhões de listings no Google Business Profile globalmente. Cadastra e gerencia presença digital.', 'Tornar seu estabelecimento visível no mapa para atrair clientes, exibindo horário, telefone e fotos.'),
        ('Ator 4 (Integrador)', 'Desenvolvedor: 1,2+ milhão de empresas usam Google Maps Platform APIs. Reclassificado de "Opcional" para "Integrador" por representar ~70% da receita do Google Maps (~$11,8 bilhões em 2023).', 'Incorporar funcionalidades de mapa em sistemas externos via API (Routes, Places, Maps SDK).'),
    ])
    criar_fonte(doc, 'Fontes: Alphabet Inc. 10-K FY2023 | Google Maps Platform Docs | Local Guides Connect')

    criar_secao_titulo(doc, '2. Liste 10 Funcionalidades Principais')
    criar_instrucao(doc, 'Funcionalidades verificadas no sistema real com detalhes técnicos.')
    criar_tabela_funcionalidades(doc, [
        (1, 'Buscar endereço, nome de estabelecimento ou coordenadas no mapa', 'Usuário'),
        (2, 'Calcular rota entre origem e destino com opções de transporte (carro, ônibus, a pé, bicicleta)', 'Usuário'),
        (3, 'Navegar em tempo real com instruções de voz e direção passo a passo via GPS', 'Usuário'),
        (4, 'Visualizar informações detalhadas de um local (horário, telefone, avaliações, fotos)', 'Usuário'),
        (5, 'Avaliar e comentar sobre um estabelecimento após visitá-lo', 'Local Guide'),
        (6, 'Cadastrar estabelecimento comercial com endereço, categoria e dados de contato', 'Comerciante'),
        (7, 'Baixar área do mapa para uso off-line (sem conexão com a internet)', 'Usuário'),
        (8, 'Compartilhar localização em tempo real com contatos selecionados', 'Usuário'),
        (9, 'Visualizar condições de trânsito em tempo real com cores no mapa (verde, amarelo, vermelho)', 'Usuário'),
        (10, 'Salvar locais favoritos (casa, trabalho, restaurantes) para acesso rápido', 'Usuário'),
    ])

    criar_secao_titulo(doc, '3. Priorize com MoSCoW')
    criar_instrucao(doc, 'Must Have ≤ 60% do esforço (DSDM). Trânsito (#9) reclassificado de Must para Should com justificativa.')
    criar_moscow(doc,
        ['#1 - Buscar endereço', '#2 - Calcular rota', '#4 - Visualizar informações do local'],
        ['#3 - Navegação GPS em tempo real', '#6 - Cadastrar estabelecimento', '#9 - Trânsito em tempo real', '#10 - Salvar locais favoritos'],
        ['#5 - Avaliar estabelecimentos', '#7 - Mapa off-line', '#8 - Compartilhar localização'],
        ['Street View (visualização 360° das ruas)', 'Reserva de restaurantes direto pelo mapa', 'Integração com assistente de voz para comandos mãos-livres'],
        nota_analitica='Nota: trânsito (#9) movido de Must para Should. A Routes API oferece modo TRAFFIC_UNAWARE como fallback. O Maps funciona sem dados de trânsito, apenas com qualidade inferior. DSDM: Must Have ≤ 60% do esforço.'
    )

    criar_secao_titulo(doc, '4. Requisitos Não Funcionais')
    criar_instrucao(doc, 'Baseados no modelo ISO 25010 e em dados reais da infraestrutura do Google Maps.')
    criar_rnf(doc, [
        ('Desempenho', 'O mapa deve carregar em até 3 segundos em conexão 4G. As sugestões de busca devem aparecer em até 500ms. A plataforma processa bilhões de requisições diárias com 2 bilhões de MAU.', 'Tempo de carregamento < 3s (4G); latência de autocomplete < 500ms.'),
        ('Segurança', 'Localização coletada somente com consentimento explícito (Geolocation API). Toda transmissão via HTTPS com TLS 1.2+. Conformidade com GDPR, LGPD e leis locais de privacidade.', '100% das conexões via HTTPS; zero coletas sem permissão; TLS 1.2+ (corrigido de 1.3, por backward compatibility com dispositivos antigos).'),
        ('Disponibilidade', 'SLA publicado de 99,9% de uptime para Google Maps Platform APIs. Em fevereiro de 2026, um incidente (ID 163VbP2njo8deNYScb4v) causou ~36h de instabilidade, gerando fallback para modo TRAFFIC_UNAWARE.', 'Uptime ≥ 99,9% (SLA contratual); monitoramento via Google Cloud Status Dashboard.'),
        ('Usabilidade', 'Operável com uma mão (uso em movimento). Diretrizes WCAG 2.1 AA. Botões mínimos de 48×48px. Interface disponível em 70+ idiomas.', 'Aprovação WCAG 2.1 nível AA; tamanho mínimo de botões: 48×48px.'),
        ('Escalabilidade', 'O sistema suporta 2 bilhões de usuários ativos mensais. Infraestrutura global do Google Cloud com data centers em 40+ regiões.', 'Baseado em 2 bilhões de MAU reais (corrigido de "100 milhões simultâneos", dado inventado na ficha original).'),
    ])
    criar_fonte(doc, 'Fontes: Google Maps Platform SLA | Google Cloud Status Dashboard | Alphabet 10-K')

    criar_secao_titulo(doc, '5. Cenário de Exceção')
    criar_instrucao(doc, 'Cenário baseado em incidente real documentado no Google Cloud Status Dashboard.')
    criar_excecao(doc,
        ['Indisponibilidade de serviço externo', 'Timeout / lentidão extrema'],
        [
            'O serviço de dados de trânsito em tempo real ficou indisponível. Incidente real: fevereiro de 2026 (ID 163VbP2njo8deNYScb4v), ~36h de instabilidade na Routes API afetando dados de trânsito em múltiplas regiões.',
            'Usuário em navegação ativa: recebe rota sem considerar trânsito. Desenvolvedores usando Routes API: recebem respostas degradadas ou timeouts. Comerciantes: dados de popularidade por horário ficam indisponíveis.',
            'A Routes API oferece modo TRAFFIC_UNAWARE como fallback automático: calcula rotas por distância e velocidade-limite, sem dados de trânsito em tempo real. O mapa off-line continua funcionando se previamente baixado. O sistema exibe aviso de degradação.',
            '"Informações de trânsito temporariamente indisponíveis. A rota sugerida pode não refletir as condições atuais."',
            'Rotas são calculadas normalmente usando dados estáticos. O cache local mantém dados recentes por período configurável. O sistema tenta reconectar automaticamente com backoff exponencial. Logs de incidente são criados para monitoramento SRE.',
        ]
    )
    criar_fonte(doc, 'Fonte: Google Cloud Status Dashboard, Incident 163VbP2njo8deNYScb4v (fev/2026)')

    criar_secao_titulo(doc, '6. Pós-Análise (Validação com o app real)')
    criar_pos_analise(doc, [
        ('Todos os atores estão corretos?', 'Sim. Ator 2 reclassificado para Colaborador; Ator 4 reclassificado para Integrador com base em impacto de receita.'),
        ('As 10 funcionalidades existem no sistema?', 'Todas verificadas. Detalhes do TRAFFIC_UNAWARE fallback confirmados na documentação da API.'),
        ('A priorização MoSCoW bate com a realidade?', '#9 movido de Must para Should, pois o sistema funciona sem trânsito via TRAFFIC_UNAWARE.'),
        ('Os requisitos não funcionais são viáveis?', 'SLA corrigido para 99,9% (publicado). TLS corrigido para 1.2+. Escalabilidade baseada em 2B MAU reais.'),
        ('O cenário de exceção é tratado pelo sistema?', 'Confirmado com incidente real de fev/2026. Fallback TRAFFIC_UNAWARE documentado na Routes API.'),
    ])

    criar_secao_titulo(doc, '7. Observações Finais / Aprendizados')
    criar_observacoes(doc, 'A análise aprofundada do Google Maps revelou que a classificação inicial de atores subestimava o papel de dois deles. O Local Guide, originalmente classificado como "Suporte", é na verdade um Colaborador, pois ele gera conteúdo que alimenta o sistema (avaliações, fotos, correções), não apenas auxilia outros usuários. O Desenvolvedor, classificado como "Opcional", é na verdade um Integrador estratégico, pois representa ~70% da receita do Google Maps (~$11,8 bilhões/ano). Essa reclassificação demonstra que a análise de impacto econômico e operacional deve complementar a classificação funcional de atores.\n\nNa priorização MoSCoW, o trânsito em tempo real (#9) foi movido de Must Have para Should Have após verificação técnica: a Routes API oferece modo TRAFFIC_UNAWARE como fallback, provando que o sistema funciona sem dados de trânsito, apenas com qualidade inferior. Essa correção reflete o princípio DSDM de que Must Have ≤ 60% do esforço.\n\nOs requisitos não funcionais foram corrigidos com dados reais: o SLA publicado é 99,9% (não 99,99%), o TLS mínimo é 1.2+ (não 1.3, por backward compatibility), e a escalabilidade é baseada em 2 bilhões de MAU reais (não "100 milhões simultâneos", que era um dado inventado). O incidente de fevereiro de 2026 (ID 163VbP2njo8deNYScb4v, ~36h de instabilidade) demonstra que mesmo sistemas com SLA de 99,9% sofrem interrupções significativas, o que reforça a importância de mapear cenários de exceção com fallbacks concretos.')

    criar_rodape(doc, 'Google Maps', 1)
    path = os.path.join(OUTPUT_DIR, 'Ficha_Requisitos_01_Google_Maps.docx')
    doc.save(path)
    print(f'Criado: {path}')


# ===========================================================================
# FICHA 2 — AIRBNB (Análise Aprofundada)
# ===========================================================================
def criar_ficha_airbnb():
    doc = Document()
    configurar_doc(doc)
    criar_cabecalho(doc, 2)
    criar_titulo(doc, 'Airbnb', 'marketplace bilateral de hospedagens | 8 milhões de anúncios ativos em 220 países')
    criar_ficha_id(doc, 'Airbnb')

    criar_secao_titulo(doc, '1. Mapeie os Atores')
    criar_instrucao(doc, 'O Airbnb é um marketplace bilateral: dois atores com interesses opostos operam sobre o mesmo recurso (o espaço físico).')
    criar_tabela_atores(doc, [
        ('Ator 1 (Principal)', 'Hóspede: viajante que busca acomodação temporária. ~1,34 milhão de reservas por dia (~491 milhões de noites em 2024).', 'Encontrar e reservar um espaço com bom custo-benefício, de forma segura e rápida.'),
        ('Ator 2 (Mediador)', 'Suporte ao Cliente: equipe com poder de decisão: processa reembolsos, media disputas e pode cancelar reservas. Reclassificado de "Suporte" para "Mediador" por executar ações com impacto financeiro direto.', 'Mediar conflitos com critérios de alçada definidos, garantindo satisfação de ambas as partes e conformidade com as políticas.'),
        ('Ator 3 (Fornecedor)', 'Anfitrião: 5+ milhões de anfitriões ativos gerenciando 8 milhões de anúncios. Penalidade por double-booking: $50–$100 + bloqueio de calendário.', 'Cadastrar espaço, gerenciar disponibilidade/preços e receber pagamentos. GBV: ~$82 bilhões/ano (2024).'),
        ('Ator 4 (Administrador)', 'Administrador da Plataforma: modera anúncios, verifica identidades, define 6 tipos de políticas de cancelamento.', 'Garantir qualidade e segurança dos anúncios, conformidade com LGPD, PSD2, PCI-DSS.'),
    ])
    criar_fonte(doc, 'Fontes: Airbnb Q4 2024 Earnings (SEC 8-K) | DemandSage Statistics 2025 | Hostify Blog')

    criar_secao_titulo(doc, '2. Liste 10 Funcionalidades Principais')
    criar_instrucao(doc, 'Em marketplace bilateral, cada funcionalidade deve ser avaliada sob a ótica de ambos os lados.')
    criar_tabela_funcionalidades(doc, [
        (1, 'Buscar acomodações por destino, datas de check-in/check-out e número de hóspedes', 'Hóspede'),
        (2, 'Cadastrar anúncio com fotos, descrição, preço por noite e regras da casa', 'Anfitrião'),
        (3, 'Reservar acomodação com pagamento on-line seguro, processado via Braintree (PCI-DSS Nível 1)', 'Hóspede'),
        (4, 'Gerenciar calendário de disponibilidade com sincronização iCal para múltiplas plataformas', 'Anfitrião'),
        (5, 'Avaliar a estadia (sistema double-blind: nenhum lado vê a avaliação do outro até ambos enviarem ou 14 dias expirarem', 'Hóspede'),
        (6, 'Avaliar o hóspede com nota e comentário (mesmo sistema double-blind)', 'Anfitrião'),
        (7, 'Filtrar resultados por preço, tipo de espaço, comodidades e avaliação mínima', 'Hóspede'),
        (8, 'Enviar e receber mensagens entre hóspede e anfitrião pelo chat interno', 'Hóspede / Anfitrião'),
        (9, 'Cancelar reserva conforme 1 das 6 políticas (Flexível, Moderada, Limitada, Firme, Rigorosa, Super Rigorosa)', 'Hóspede / Anfitrião'),
        (10, 'Visualizar localização do imóvel no mapa com pontos de interesse próximos', 'Hóspede'),
    ])
    criar_fonte(doc, 'Fontes: Airbnb Help Center | Uplisting Review Policy Guide | Zeevou Cancellation Policy Guide')

    criar_secao_titulo(doc, '3. Priorize com MoSCoW')
    criar_instrucao(doc, 'Must Have ≤ 60% do esforço (DSDM).')
    criar_moscow(doc,
        ['#1 - Buscar acomodações', '#2 - Cadastrar anúncio', '#3 - Reservar com pagamento', '#4 - Gerenciar calendário'],
        ['#7 - Filtrar resultados', '#8 - Chat interno', '#9 - Cancelar reserva'],
        ['#5 - Avaliar estadia (double-blind)', '#6 - Avaliar hóspede', '#10 - Mapa com localização'],
        ['Experiências Airbnb (atividades turísticas guiadas)', 'Programa de fidelidade com pontos', 'Seguro-viagem integrado'],
        nota_analitica='Nota: calendário (#4) é Must Have porque sem ele o anfitrião não controla disponibilidade, inviabilizando toda a operação. Funcionalidades de bastidor podem ser tão críticas quanto as de front-end.'
    )

    criar_secao_titulo(doc, '4. Requisitos Não Funcionais')
    criar_instrucao(doc, 'Baseados no modelo ISO 25010 e em dados reais da infraestrutura do Airbnb.')
    criar_rnf(doc, [
        ('Desempenho', 'Busca deve retornar em até 3s. O sistema processa ~1,34 milhão de reservas/dia (~17 escritas/s), com leituras estimadas em ~1.000 QPS, fortemente cacheadas.', 'Tempo de resposta da busca < 3s (P95); latência da API < 500ms.'),
        ('Segurança', 'Pagamentos via Braintree (PCI-DSS Nível 1). Airbnb nunca armazena números de cartão (tokenização delegada). Framework Orpheus garante 99,999% de consistência via chaves de idempotência e DAGs retentáveis.', 'Certificação PCI-DSS ativa via processador; zero PANs em texto; conformidade LGPD, GDPR, PSD2/PSD3.'),
        ('Disponibilidade', 'Sem SLA formal publicado. AWS Multi-AZ (99,99% garantido pela AWS). Monitoramento indica ~8h40min por incidente e ~3,4 incidentes/mês.', 'Meta implícita ≥ 99,9% mensal; monitoramento contínuo via IsDown/StatusGator.'),
        ('Usabilidade', 'Fluxo de reserva em máximo 4 cliques. Mobile First. 10+ idiomas. Frontend usa chave de idempotência (reservationId) contra duplo clique.', 'Máximo 4 cliques até confirmação; teste em dispositivos de 5" a 13".'),
        ('Escalabilidade', '200 milhões de usuários ativos, 8 milhões de anúncios, 5–10 check-ins/s contínuos. Arquitetura evoluiu de monolito Rails para ~2.000 microsserviços (SOA v2 com Service Blocks).', 'Kubernetes + AWS; database-per-service (PostgreSQL + MySQL/RDS); Kafka para eventos assíncronos.'),
    ])
    criar_fonte(doc, 'Fontes: Airbnb Tech Blog (Orpheus) | SEC 10-K FY2024 | ByteByteGo Architecture | Braintree PCI')

    criar_secao_titulo(doc, '5. Cenário de Exceção')
    criar_instrucao(doc, 'Cenário baseado em arquitetura real documentada pelo Airbnb Engineering Blog.')
    criar_excecao(doc,
        ['Falha de pagamento', 'Conflito de agendamento'],
        [
            'Dois hóspedes tentam reservar o mesmo imóvel nas mesmas datas simultaneamente. A ~17 transações de escrita por segundo, esse cenário é estatisticamente frequente. Há também risco de double-booking cross-platform (Airbnb + Booking.com + Vrbo).',
            'Segundo hóspede: reserva recusada. Anfitrião: pode receber notificações contraditórias. No caso cross-platform, penalidade de $50–$100 e bloqueio de calendário.',
            'O Airbnb usa transações ACID com PostgreSQL e lock de linha (row-level locking), mantendo reserva e inventário no mesmo banco para evitar inconsistência. Decisão deliberada de NÃO usar saga/two-phase commit (consistência eventual inaceitável para reservas). Frontend envia chave de idempotência (reservationId) contra duplo clique.',
            '"Este espaço acabou de ser reservado para as datas selecionadas. Veja datas disponíveis ou acomodações semelhantes na região."',
            'Reserva do primeiro confirmada e datas bloqueadas via lock de linha no PostgreSQL. Nenhuma cobrança ao segundo. Framework Orpheus converte entrega at-least-once do Kafka em garantia exactly-once para pagamentos via DAGs idempotentes. Para cross-platform, protocolo iCal pode levar 20min a horas; solução profissional é channel manager com API direta.',
        ]
    )
    criar_fonte(doc, 'Fontes: Airbnb Tech Blog, "Avoiding Double Payments" | ITNEXT, "Solving Double Booking at Scale"')

    criar_secao_titulo(doc, '6. Pós-Análise (Validação com o app real)')
    criar_pos_analise(doc, [
        ('Todos os atores estão corretos?', 'Sim. Suporte reclassificado como Mediador, com poder de decisão financeira.'),
        ('As 10 funcionalidades existem no sistema?', 'Todas verificadas. Double-blind e 6 políticas de cancelamento confirmados.'),
        ('A priorização MoSCoW bate com a realidade?', 'Calendário como Must Have validado, pois sem ele não há operação.'),
        ('Os requisitos não funcionais são viáveis?', 'PCI-DSS via Braintree confirmado. Orpheus 99,999% confirmado no Tech Blog.'),
        ('O cenário de exceção é tratado pelo sistema?', 'Row-level locking com PostgreSQL confirmado. Chave de idempotência no frontend confirmada.'),
    ])

    criar_secao_titulo(doc, '7. Observações Finais / Aprendizados')
    criar_observacoes(doc, 'A análise do Airbnb introduz um aspecto que o Google Maps não possui: a existência de dois atores com interesses opostos operando sobre o mesmo recurso. O hóspede quer disponibilidade e preço baixo; o anfitrião quer ocupação e preço alto. Essa dinâmica de marketplace bilateral tem implicações diretas na modelagem de requisitos.\n\nO mapeamento de atores revelou que o Suporte ao Cliente é um ator com poder de decisão financeira (reembolsos, cancelamentos), exigindo requisitos de segurança e auditoria específicos. Na priorização MoSCoW, o calendário (#4) como Must Have demonstra que funcionalidades de bastidor podem ser tão críticas quanto as de front-end.\n\nOs requisitos não funcionais trouxeram PCI-DSS (padrão obrigatório para processamento de cartões) e o framework Orpheus (99,999% de consistência em pagamentos via DAGs idempotentes). O cenário de exceção revelou a decisão arquitetural de usar ACID com row-level locking em PostgreSQL, recusando padrões distribuídos (saga, two-phase commit) porque consistência eventual é inaceitável quando dinheiro está envolvido.')

    criar_rodape(doc, 'Airbnb', 2)
    path = os.path.join(OUTPUT_DIR, 'Ficha_Requisitos_02_Airbnb.docx')
    doc.save(path)
    print(f'Criado: {path}')


# ===========================================================================
# FICHA 3 — UBER (Análise Aprofundada)
# ===========================================================================
def criar_ficha_uber():
    doc = Document()
    configurar_doc(doc)
    criar_cabecalho(doc, 3)
    criar_titulo(doc, 'Uber', 'plataforma de mobilidade sob demanda | 36 milhões de viagens/dia em 70+ países')
    criar_ficha_id(doc, 'Uber')

    criar_secao_titulo(doc, '1. Mapeie os Atores')
    criar_instrucao(doc, 'O Uber opera um marketplace bilateral de tempo real, onde a janela de decisão é de segundos, não dias.')
    criar_tabela_atores(doc, [
        ('Ator 1 (Principal)', 'Passageiro, com ~161 milhões de MAU (Q4 2024). Solicita viagem em tempo real.', 'Solicitar e concluir uma viagem do ponto A ao ponto B com segurança, previsibilidade de preço e tempo de espera curto.'),
        ('Ator 2 (Fornecedor)', 'Motorista parceiro, com ~9,7 milhões ativos globalmente. Classificados como contratantes independentes. Decisão de aceite deve ocorrer em <10s.', 'Aceitar corridas, completar viagens e maximizar ganhos.'),
        ('Ator 3 (Sistema)', 'Algoritmo de Matching: sistema automatizado com grafo bipartido e janelas de ~10s. Otimiza múltiplos fatores simultaneamente (Hungarian Algorithm adaptado).', 'Minimizar tempo de espera do passageiro E tempo ocioso do motorista simultaneamente.'),
        ('Ator 4 (Administrador)', 'Administrador da Plataforma: gerencia surge pricing, segurança, verificação de documentos e conformidade regulatória em 70+ países.', 'Garantir operação segura e regulamentada, calibrar algoritmos de pricing e matching.'),
    ])
    criar_fonte(doc, 'Fontes: Uber SEC 10-K FY2024 | DemandSage Uber Statistics 2025 | Uber Engineering Blog')

    criar_secao_titulo(doc, '2. Liste 10 Funcionalidades Principais')
    criar_instrucao(doc, 'Em sistema de tempo real, a latência de cada funcionalidade é um requisito implícito.')
    criar_tabela_funcionalidades(doc, [
        (1, 'Solicitar viagem informando origem e destino, com estimativa de preço e tempo em tempo real', 'Passageiro'),
        (2, 'Parear passageiro com motorista via algoritmo de grafo bipartido com janelas de ~10s (Hungarian Algorithm)', 'Sistema (Matching)'),
        (3, 'Rastrear posição do motorista em tempo real, com GPS adaptativo com Filtro de Kalman e Viterbi para map-matching (<200ms)', 'Passageiro / Motorista'),
        (4, 'Calcular preço dinâmico (surge) por região usando grid hexagonal H3 com recálculo a cada 1–2 minutos', 'Sistema (Pricing)'),
        (5, 'Aceitar ou recusar corrida com informações de destino, valor e distância', 'Motorista'),
        (6, 'Avaliar motorista com nota de 1 a 5 estrelas ao final da corrida', 'Passageiro'),
        (7, 'Avaliar passageiro com nota de 1 a 5 estrelas ao final da corrida', 'Motorista'),
        (8, 'Processar pagamento automático com split instantâneo (~25% Uber vs. ~75% motorista)', 'Sistema (Pagamento)'),
        (9, 'Compartilhar localização e status da viagem com contatos de confiança', 'Passageiro'),
        (10, 'Navegar até o destino com GPS turn-by-turn, usando provedores híbridos de mapa', 'Motorista'),
    ])

    criar_secao_titulo(doc, '3. Priorize com MoSCoW')
    criar_instrucao(doc, 'Must Have ≤ 60% do esforço (DSDM). Em tempo real, matching e tracking são infraestrutura.')
    criar_moscow(doc,
        ['#1 - Solicitar viagem', '#2 - Matching passageiro-motorista', '#3 - Rastreamento GPS', '#5 - Aceitar/recusar corrida', '#8 - Processar pagamento'],
        ['#4 - Preço dinâmico (surge)', '#10 - Navegação GPS integrada'],
        ['#6 - Avaliar motorista', '#7 - Avaliar passageiro', '#9 - Compartilhar viagem'],
        ['Uber Eats (entrega de comida)', 'Uber Reserve (agendamento antecipado)', 'Uber Teen (contas para menores)'],
        nota_analitica='Nota: matching (#2) é Must Have, pois é o núcleo da proposta de valor. Surge (#4) ficou em Should Have: a Uber operou inicialmente com preço fixo, mas o surge é essencial para equilibrar oferta/demanda em escala.'
    )

    criar_secao_titulo(doc, '4. Requisitos Não Funcionais')
    criar_instrucao(doc, 'Baseados no modelo ISO 25010 e em dados reais da infraestrutura da Uber.')
    criar_rnf(doc, [
        ('Desempenho', 'Matching em <10s. GPS com <200ms de latência usando Filtro de Kalman (suavização) e Viterbi (map-matching). Milhões de atualizações de localização processadas por segundo globalmente.', 'Matching P95 < 10s; GPS latency P99 < 200ms; throughput > 1M location updates/s.'),
        ('Segurança', 'MFA para motoristas e passageiros. Em setembro 2022, brecha via "MFA Fatigue" (grupo Lapsus$). Atacante bombardeou funcionário com notificações MFA. Desde então: migração para FIDO2 e zero-trust para acesso interno.', '100% motoristas com verificação de identidade; MFA obrigatório; zero acessos admin sem FIDO2.'),
        ('Disponibilidade', 'Sem SLA formal publicado. Arquitetura multi-região ativa-ativa em data centers próprios + Google Cloud + Oracle Cloud. Cada minuto de downtime afeta ~25.000 viagens em andamento.', 'Uptime ≥ 99,99%; failover entre regiões < 30s.'),
        ('Usabilidade', 'Solicitar viagem em máximo 3 toques. Interface otimizada para uso com uma mão em movimento. Mapa deve carregar em <2s mesmo em 3G.', 'Máximo 3 toques até solicitação; FCP < 2s em 3G; suporte VoiceOver/TalkBack.'),
        ('Escalabilidade', '36 milhões de viagens/dia com 9,7 milhões de motoristas. Arquitetura DOMA: ~2.200 microsserviços em 70 domínios com interfaces de fachada. Geolocalização via H3 (grid hexagonal open-source criado pela Uber).', 'Picos de 2x demanda normal sem degradação; auto-scaling por domínio; H3 resolution 7 (~5,16 km²) para surge.'),
    ])
    criar_fonte(doc, 'Fontes: Uber Engineering Blog (DOMA, H3, Kalman) | SEC 10-K FY2024 | CSO Online, 2022 Breach')

    criar_secao_titulo(doc, '5. Cenário de Exceção')
    criar_instrucao(doc, 'Cenário baseado em comportamento real documentado pela engenharia da Uber.')
    criar_excecao(doc,
        ['Falha de GPS / perda de sinal', 'Nenhum motorista disponível'],
        [
            'Passageiro solicita viagem mas não há motoristas disponíveis na região dentro da janela de matching (~10s). Em paralelo, GPS retorna coordenadas imprecisas (drift em áreas urbanas com prédios altos).',
            'Passageiro: não consegue viagem. Motoristas próximos: podem perder janela de matching. Sistema de surge pode interpretar baixa oferta como sinal para aumentar preços.',
            '1) Matching expande raio progressivamente (1km → 3km → 5km). 2) Ativa surge para atrair motoristas de regiões adjacentes (preço como redistribuição geográfica). 3) Filtro de Kalman suaviza GPS e Viterbi faz snap-to-road. 4) Se GPS perdido completamente: última posição + interpolação por velocidade/direção.',
            '"Nenhum motorista disponível no momento. Tentaremos novamente. Ativar notificação quando um motorista estiver disponível?" + opções alternativas (UberX vs. Comfort).',
            'Solicitação sem match registrada como "unfulfilled demand" para modelos de ML de predição. GPS corrigido armazenado com confidence score. Nenhuma cobrança gerada. Regiões com déficit crônico alimentam decisões de incentivo (bônus por zona/horário).',
        ]
    )

    criar_secao_titulo(doc, '6. Pós-Análise (Validação com o app real)')
    criar_pos_analise(doc, [
        ('Todos os atores estão corretos?', 'Sim. Algoritmo de Matching incluído como ator-sistema, com autonomia de decisão.'),
        ('As 10 funcionalidades existem no sistema?', 'Todas verificadas. H3, Kalman e DOMA confirmados no Engineering Blog.'),
        ('A priorização MoSCoW bate com a realidade?', 'Surge como Should Have validado, pois a Uber operou sem ele inicialmente.'),
        ('Os requisitos não funcionais são viáveis?', 'Breach de 2022 documentado. DOMA e H3 open-source confirmados.'),
        ('O cenário de exceção é tratado pelo sistema?', 'Expansão de raio e surge como redistribuição geográfica confirmados.'),
    ])

    criar_secao_titulo(doc, '7. Observações Finais / Aprendizados')
    criar_observacoes(doc, 'A análise do Uber introduz o tempo real como requisito arquitetural fundante, e a janela de ~10 segundos para matching define a arquitetura inteira. O algoritmo de matching é um ator-sistema com autonomia de decisão (grafo bipartido, Hungarian Algorithm), cujos critérios devem ser documentados como requisitos auditáveis.\n\nO matching acumula solicitações em janelas de ~10s para otimizar emparelhamento global, um trade-off (latência individual vs. eficiência global) que não aparece em entrevistas com stakeholders e exige análise técnica independente.\n\nO breach de 2022 (Lapsus$, MFA Fatigue) demonstra que requisitos de segurança incluem engenharia social e fadiga humana. A arquitetura DOMA organiza ~2.200 microsserviços em 70 domínios com interfaces de fachada. O H3 (grid hexagonal open-source) divide o planeta em hexágonos uniformes. A escolha por hexágonos elimina viés direcional em cálculos de proximidade, com recálculo de surge a cada 1–2 minutos.')

    criar_rodape(doc, 'Uber', 3)
    path = os.path.join(OUTPUT_DIR, 'Ficha_Requisitos_03_Uber.docx')
    doc.save(path)
    print(f'Criado: {path}')


# ===========================================================================
# FICHA 4 — TRELLO (Análise Aprofundada)
# ===========================================================================
def criar_ficha_trello():
    doc = Document()
    configurar_doc(doc)
    criar_cabecalho(doc, 4)
    criar_titulo(doc, 'Trello', 'ferramenta de gestão visual Kanban | 50+ milhões de usuários registrados')
    criar_ficha_id(doc, 'Trello')

    criar_secao_titulo(doc, '1. Mapeie os Atores')
    criar_instrucao(doc, 'O Trello é um sistema colaborativo, onde múltiplos atores operam simultaneamente sobre o mesmo recurso (o quadro).')
    criar_tabela_atores(doc, [
        ('Ator 1 (Principal)', 'Membro do Quadro: cria, move e edita cartões. 50+ milhões de usuários. Plano free: limite de 10 colaboradores por quadro.', 'Organizar tarefas visualmente em colunas (listas), movendo cartões entre etapas do fluxo de trabalho.'),
        ('Ator 2 (Administrador)', 'Administrador do Workspace: 3 níveis de papel no workspace (admin, normal, guest) e 3 no quadro (admin, normal, observer).', 'Controlar acesso, gerenciar licenças (Free, Standard, Premium, Enterprise) e políticas de visibilidade.'),
        ('Ator 3 (Sistema)', 'Power-Ups / Butler: integrações e automações. Butler permite regras, botões e comandos de calendário. Power-Ups conectam Slack, GitHub, Drive.', 'Automatizar ações repetitivas e integrar dados externos.'),
        ('Ator 4 (Observador)', 'Observer: papel somente leitura (Premium/Enterprise). Pode visualizar quadros mas não editar.', 'Acompanhar progresso sem risco de alteração acidental. Para stakeholders e gestores.'),
    ])
    criar_fonte(doc, 'Fontes: Trello Help, Board Permissions | Atlassian Pricing Page | Business of Apps 2025')

    criar_secao_titulo(doc, '2. Liste 10 Funcionalidades Principais')
    criar_instrucao(doc, 'Modelo mental: quadro → listas (colunas) → cartões (itens de trabalho).')
    criar_tabela_funcionalidades(doc, [
        (1, 'Criar quadro com listas (colunas) personalizáveis representando etapas do fluxo de trabalho', 'Membro / Admin'),
        (2, 'Criar, editar e excluir cartões com título, descrição (Markdown), membros, etiquetas e data de entrega', 'Membro'),
        (3, 'Arrastar e soltar cartões entre listas, sincronizado em tempo real via WebSocket (Socket.io modificado)', 'Membro'),
        (4, 'Adicionar checklists com itens marcáveis e barra de progresso percentual', 'Membro'),
        (5, 'Anexar arquivos ao cartão (upload direto, Google Drive, Dropbox). 10MB/arquivo free, 250MB pagos', 'Membro'),
        (6, 'Comentar em cartões com menções (@membro) que geram notificação por e-mail e push', 'Membro'),
        (7, 'Filtrar cartões por etiqueta, membro, data de entrega ou texto', 'Membro'),
        (8, 'Configurar automações via Butler: regras (quando X, faça Y), botões e comandos agendados', 'Membro / Admin'),
        (9, 'Gerenciar permissões de membros no quadro (admin, normal, observer) e visibilidade', 'Admin do Workspace'),
        (10, 'Visualizar atividade do quadro em feed cronológico (Activity Log) com histórico completo', 'Membro / Observer'),
    ])

    criar_secao_titulo(doc, '3. Priorize com MoSCoW')
    criar_instrucao(doc, 'Must Have ≤ 60% do esforço (DSDM). O Trello é intencionalmente simples, e o valor está na restrição.')
    criar_moscow(doc,
        ['#1 - Criar quadro com listas', '#2 - Criar/editar cartões', '#3 - Drag-and-drop com sync em tempo real', '#9 - Gerenciar permissões'],
        ['#4 - Checklists', '#6 - Comentários com @menções', '#7 - Filtrar cartões'],
        ['#5 - Anexar arquivos', '#8 - Automações Butler', '#10 - Activity Log'],
        ['Visualização Timeline (Gantt)', 'Dashboard com métricas de produtividade', 'Integração nativa com Jira (bridge bidirecional)'],
        nota_analitica='Nota: drag-and-drop com sync (#3) é Must Have, pois é o diferencial fundante do Trello. Butler (#8) como Could Have reflete a história: o Trello operou por anos sem ele (adquirido e integrado apenas em 2018).'
    )

    criar_secao_titulo(doc, '4. Requisitos Não Funcionais')
    criar_instrucao(doc, 'Baseados no modelo ISO 25010 e na stack tecnológica real do Trello.')
    criar_rnf(doc, [
        ('Desempenho', 'Sync de drag-and-drop em <500ms para todos os clientes. WebSocket via Socket.io modificado, com fallback para HTTP long-polling. Limite de 5.000 cartões abertos por quadro (hard limit), e a performance degrada a partir de ~1.000 por lista.', 'Latência sync P95 < 500ms; abertura de quadro com 500 cartões < 3s; limite 5.000 cartões/quadro.'),
        ('Segurança', 'TLS 1.2+. AES-256 em repouso. SSO (SAML 2.0) para Premium/Enterprise. 2FA para todos. Conformidade SOC 2 Type II e ISO 27001 (herdada da Atlassian Cloud).', 'SOC 2 e ISO 27001 ativas; zero dados em texto puro em logs; sessões com expiração configurável.'),
        ('Disponibilidade', 'Atlassian Cloud SLA: 99,9% para Premium/Enterprise. Stack: Node.js + MongoDB + Redis + HAProxy na AWS. Incidentes tipicamente resolvidos em <4h.', 'Uptime ≥ 99,9% mensal (≤ 43 min downtime/mês); MTTR < 4h para incidentes críticos.'),
        ('Usabilidade / Acessibilidade', 'Mobile First responsivo. Modo colorblind com padrões de textura nas etiquetas desde 2014 (WCAG 2.1 AA). Navegação por teclado para todas as operações.', 'WCAG 2.1 AA; modo colorblind ativável; todas ações via teclado (Tab, Enter, Esc).'),
        ('Escalabilidade', '50+ milhões de usuários. MongoDB (escolha natural para cartões com campos variáveis, o que evita múltiplas tabelas de junção do modelo relacional). Redis Cluster para WebSocket sessions. Auto-scaling Node.js na AWS.', 'Sharding MongoDB; Redis Cluster; auto-scaling Node.js.'),
    ])
    criar_fonte(doc, 'Fontes: Trello Engineering Blog | Atlassian Trust Center (SOC 2, ISO 27001) | Atlassian Cloud SLA')

    criar_secao_titulo(doc, '5. Cenário de Exceção')
    criar_instrucao(doc, 'Cenário baseado na estratégia real de resolução de conflitos do Trello.')
    criar_excecao(doc,
        ['Edição simultânea / conflito de dados', 'Perda de conexão WebSocket'],
        [
            'Dois membros editam o mesmo cartão simultaneamente: um altera o título, outro modifica a descrição. Um terceiro arrasta o cartão para outra lista. As três operações chegam ao servidor em ordem diferente da executada.',
            'Todos os membros do quadro, que podem ver estados inconsistentes por frações de segundo. O membro cuja edição chegou por último tem sua versão preservada (last-write-wins).',
            'Trello usa server-authoritative last-write-wins: servidor é fonte de verdade, última operação prevalece. NÃO usa CRDTs nem Operational Transformation, pois essas técnicas são para editores de texto (Google Docs), mas o modelo de campos discretos do Trello (título, descrição, posição) torna last-write-wins suficiente. Cada campo é atualizado independentemente.',
            'Sem mensagem explícita de conflito. A interface atualiza via WebSocket. Se conexão cair: fallback HTTP long-polling + indicador "Reconectando...". Ao reconectar, cliente solicita estado completo do quadro.',
            'Todas as operações registradas no Activity Log com timestamp e autor. Sem "undo" global. O membro que perdeu edição refaz manualmente. Operações destrutivas usam soft-delete (arquivo restaurável). Last-write-wins é trade-off consciente: simplicidade vs. risco de perda silenciosa, mas adequado para o volume típico do Trello (dezenas de ops/minuto, não milhares).',
        ]
    )

    criar_secao_titulo(doc, '6. Pós-Análise (Validação com o app real)')
    criar_pos_analise(doc, [
        ('Todos os atores estão corretos?', 'Sim. Observer como ator distinto, com papel exclusivo Premium/Enterprise.'),
        ('As 10 funcionalidades existem no sistema?', 'Todas verificadas. 5.000 cartões e 10MB/arquivo confirmados na documentação.'),
        ('A priorização MoSCoW bate com a realidade?', 'Butler como Could Have validado, pois foi integrado apenas em 2018.'),
        ('Os requisitos não funcionais são viáveis?', 'Stack Node.js + MongoDB + Redis confirmada. SLA 99,9% verificado.'),
        ('O cenário de exceção é tratado pelo sistema?', 'Last-write-wins confirmado, sem CRDTs/OT. Fallback HTTP long-polling verificado.'),
    ])

    criar_secao_titulo(doc, '7. Observações Finais / Aprendizados')
    criar_observacoes(doc, 'A análise do Trello introduz colaboração em tempo real sobre dados compartilhados persistentes, problema distinto dos anteriores. A estratégia last-write-wins é deliberadamente simples: sem CRDTs nem OT, porque o modelo de campos discretos (título, descrição, posição) torna essas técnicas desnecessárias. Cada campo é atualizado independentemente, e editar título não conflita com editar descrição.\n\nO Observer como ator distinto (somente leitura, Premium/Enterprise) demonstra que permissões granulares criam interseção entre requisitos funcionais e modelo de negócio. A escolha de MongoDB reflete o modelo de dados: cartões com campos variáveis (checklists, anexos, campos custom) mapeariam em múltiplas tabelas de junção no modelo relacional.\n\nO limite de 5.000 cartões (com degradação a ~1.000) emerge da implementação, sem virtualização de lista no DOM. Na especificação de requisitos, limites operacionais devem ser documentados explicitamente. A acessibilidade para daltonismo (texturas nas etiquetas desde 2014) demonstra design inclusivo que vai além da conformidade legal, oferecendo canal redundante (cor + textura) que elimina dependência exclusiva de percepção cromática.')

    criar_rodape(doc, 'Trello', 4)
    path = os.path.join(OUTPUT_DIR, 'Ficha_Requisitos_04_Trello.docx')
    doc.save(path)
    print(f'Criado: {path}')


if __name__ == '__main__':
    criar_ficha_google_maps()
    criar_ficha_airbnb()
    criar_ficha_uber()
    criar_ficha_trello()
    print('\nTodas as 4 fichas aprofundadas em Word foram criadas com sucesso!')
